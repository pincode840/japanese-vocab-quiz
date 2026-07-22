from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

path = Path("JLPT_N5_N4_통합_단어장_7-24일차.docx")
doc = Document(path)
all_text = "\n".join(
    cell.text
    for table in doc.tables
    for row in table.rows
    for cell in row.cells
)
page_breaks = sum(
    1
    for paragraph in doc.paragraphs
    for br in paragraph._p.iter(qn("w:br"))
    if br.get(qn("w:type")) == "page"
)
bad = ["すずいい", "つまら ない", "うと", "= 맞다"]
print("FILE_BYTES", path.stat().st_size)
print("TABLES", len(doc.tables))
print("DAY_TABLE_ROWS", [len(table.rows) for table in doc.tables[1:]])
print("PAGE_BREAKS", page_breaks)
print("VOCAB_ROWS", sum(len(table.rows) - 1 for table in doc.tables[1:]))
print("BAD_REMAINS", [value for value in bad if value in all_text])
print("CHECKS", all(len(table.rows) == 21 for table in doc.tables[1:]) and page_breaks == 18)
