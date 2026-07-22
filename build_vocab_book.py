import json
import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
DATA_PATH = BASE / "vocab_raw.json"
OUTPUT_PATH = BASE / "JLPT_N5_N4_통합_단어장_7-30일차.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "606A78"
HEADER_FILL = "E8EEF5"
ALT_FILL = "F7F9FC"
BORDER = "B8C6D8"
WHITE = "FFFFFF"

FONT_LATIN = "Calibri"
FONT_CJK = "Malgun Gothic"


CORRECTIONS = {
    (7, "20"): {"reading": "ばんごう"},
    (9, "3"): {"reading": "まいとし／まいねん"},
    (10, "17"): {"word": "零（0）", "meaning": "숫자 0, 영"},
    (11, "9"): {"word": "忙しい", "meaning": "바쁘다"},
    (11, "15"): {"meaning": "재미있다"},
    (12, "3"): {"reading": "すずしい", "meaning": "선선하다, 시원하다"},
    (12, "5"): {"meaning": "높다, 비싸다, 키가 크다"},
    (12, "9"): {"word": "つまらない", "meaning": "재미없다, 시시하다"},
    (12, "10"): {"meaning": "차갑다"},
    (12, "16"): {"meaning": "원하다, 갖고 싶다"},
    (12, "17"): {"meaning": "가늘다, 좁다"},
    (13, "5"): {"meaning": "싫어하다"},
    (13, "6"): {"meaning": "예쁘다, 깨끗하다"},
    (13, "7"): {"meaning": "조용하다"},
    (13, "8"): {"meaning": "잘하다, 능숙하다"},
    (13, "9"): {"meaning": "튼튼하다, 건강하다"},
    (13, "10"): {"meaning": "좋아하다"},
    (13, "11"): {"meaning": "괜찮다, 문제없다"},
    (13, "12"): {"meaning": "매우 좋아하다"},
    (13, "13"): {"meaning": "소중하다, 중요하다"},
    (13, "14"): {"meaning": "힘들다, 큰일이다"},
    (13, "15"): {"meaning": "번화하다, 활기차다"},
    (13, "17"): {"meaning": "서툴다, 잘하지 못하다"},
    (13, "18"): {"meaning": "편리하다"},
    (13, "19"): {"meaning": "유명하다"},
    (13, "20"): {"meaning": "훌륭하다"},
    (14, "2"): {"word": "開く／開く", "reading": "あく／ひらく", "meaning": "열리다／열다, 펼치다"},
    (14, "5"): {"meaning": "씻다"},
    (14, "6"): {"meaning": "있다(무생물)"},
    (14, "10"): {"meaning": "있다(생물)"},
    (14, "12"): {"meaning": "노래하다"},
    (14, "13"): {"meaning": "태어나다"},
    (14, "15"): {"meaning": "일어나다"},
    (14, "17"): {"meaning": "가르치다, 알려 주다"},
    (14, "18"): {"meaning": "외우다, 기억하다"},
    (14, "19"): {"meaning": "헤엄치다, 수영하다"},
    (14, "20"): {"meaning": "끝나다"},
    (15, "2"): {"meaning": "돌아가다"},
    (15, "4"): {"meaning": "걸다(전화를 걸다)"},
    (15, "5"): {"meaning": "빌려주다"},
    (15, "6"): {"meaning": "쓰다(모자 등을)"},
    (15, "11"): {"meaning": "대답하다"},
    (15, "12"): {"meaning": "곤란하다"},
    (16, "2"): {"meaning": "부탁하다"},
    (16, "5"): {"meaning": "사용하다"},
    (16, "6"): {"meaning": "피곤하다, 지치다"},
    (16, "7"): {"meaning": "도착하다"},
    (16, "9"): {"meaning": "근무하다"},
    (16, "10"): {"meaning": "외출하다, 나가다"},
    (16, "14"): {"meaning": "멈추다, 서다"},
    (16, "18"): {"meaning": "줄을 서다"},
    (16, "19"): {"meaning": "늘어놓다, 줄 세우다"},
    (16, "20"): {"meaning": "자다"},
    (17, "4"): {"meaning": "들어가다"},
    (17, "5"): {"meaning": "시작되다"},
    (17, "8"): {"meaning": "이야기하다, 말하다"},
    (17, "9"): {"meaning": "날씨가 개다, 맑아지다"},
    (17, "10"): {"meaning": "당기다, 끌다"},
    (17, "11"): {"meaning": "연주하다"},
    (17, "13"): {"meaning": "내리다(비·눈이)"},
    (17, "14"): {"meaning": "기다리다"},
    (17, "15"): {"meaning": "보여 주다"},
    (17, "20"): {"meaning": "잊다"},
    (20, "5"): {"meaning": "갓난아기"},
    (20, "18"): {"meaning": "접수, 접수처"},
    (20, "19"): {"reading": "うそ", "meaning": "거짓말"},
    (21, "14"): {"meaning": "춤, 무용"},
    (21, "15"): {"meaning": "부탁, 소원"},
    (22, "2"): {"meaning": "행사장, 회의장"},
    (22, "15"): {"word": "技術", "meaning": "기술"},
    (25, "10"): {"meaning": "모래"},
    (26, "10"): {"reading": "ちゅうしゃじょう"},
    (26, "16"): {"reading": "てんきよほう", "meaning": "일기예보"},
    (27, "10"): {"meaning": "교통수단"},
    (27, "18"): {"meaning": "프로그램"},
    (28, "4"): {"meaning": "수염"},
    (29, "2"): {"meaning": "한가운데"},
    (29, "8"): {"meaning": "더운물, 목욕물"},
    (29, "18"): {"meaning": "역사"},
    (30, "4"): {"meaning": "맛있다, 잘한다"},
    (30, "9"): {"meaning": "딱딱하다"},
    (30, "10"): {"meaning": "상관없다"},
    (30, "11"): {"meaning": "엄하다, 엄격하다"},
    (30, "14"): {"meaning": "굉장하다, 대단하다"},
    (30, "15"): {"meaning": "멋있다, 훌륭하다"},
    (30, "16"): {"meaning": "미지근하다"},
    (30, "19"): {"meaning": "심하다, 지독하다"},
}


def normalize_value(text):
    text = re.sub(r"\s+", " ", text.strip())
    text = re.sub(r"\s*,\s*", ", ", text)
    return text


def load_records():
    records = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    for record in records:
        key = (record["day"], str(record["no"]))
        record.update(CORRECTIONS.get(key, {}))
        for field in ("word", "reading", "meaning"):
            record[field] = normalize_value(record[field])
    return records


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT_LATIN
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CJK)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, separate, end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, DARK_BLUE),
    ):
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("JLPT N5·N4 일본어 단어장  |  7–30일차")
    set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("—  ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(p, "PAGE")
    run = p.add_run("  —")
    set_run_font(run, size=8.5, color=MUTED)


def add_centered_para(doc, text, size, color=INK, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_cover(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(66)

    add_centered_para(doc, "JLPT N5 · N4", 11, color=BLUE, bold=True, after=16)
    add_centered_para(doc, "일본어 통합 단어장", 29, color=INK, bold=True, after=8)
    add_centered_para(doc, "능력단어 7–30일차 · 총 480단어", 14, color=DARK_BLUE, after=34)
    add_centered_para(doc, "하루 20개씩, 세 번 확인하며 완성하는 복습형 단어장", 10.5, color=MUTED, italic=True, after=42)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(10)
    run = heading.add_run("이렇게 공부하세요")
    set_run_font(run, size=14, color=BLUE, bold=True)

    steps = [
        ("1회", "일본어 단어를 보고 소리 내어 읽은 뒤 뜻을 확인합니다."),
        ("2회", "뜻을 가리고 읽기와 의미를 스스로 떠올립니다."),
        ("3회", "한국어 뜻만 보고 일본어 단어를 말하거나 써 봅니다."),
    ]
    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, [1080, 8280], indent_dxa=140)
    set_table_borders(table, color="D6DEE8", size="4")
    for i, (label, detail) in enumerate(steps):
        row = table.rows[i]
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_margins(cell, top=110, bottom=110, start=140, end=140)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(row.cells[0], HEADER_FILL)
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, size=10, bold=True, color=DARK_BLUE)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(detail)
        set_run_font(run, size=10, color=INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("복습 체크  □ 1회   □ 2회   □ 3회")
    set_run_font(run, size=10, color=MUTED, bold=True)


def add_day_page(doc, day, records):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"DAY {day}")
    set_run_font(run, size=10, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"{day}일차 핵심 단어 20")
    set_run_font(run, size=18, color=INK, bold=True)

    table = doc.add_table(rows=1, cols=5)
    widths = [540, 2160, 2340, 3060, 1260]
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    labels = ["No.", "단어", "읽기", "뜻", "복습"]
    for cell, label in zip(header.cells, labels):
        shade_cell(cell, HEADER_FILL)
        set_cell_margins(cell, top=100, bottom=100, start=120, end=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)

    for idx, item in enumerate(records, start=1):
        cells = table.add_row().cells
        row = table.rows[-1]
        prevent_row_split(row)
        if idx % 2 == 0:
            for cell in cells:
                shade_cell(cell, ALT_FILL)
        values = [str(idx), item["word"], item["reading"], item["meaning"], "□  □  □"]
        for col, (cell, value) in enumerate(zip(cells, values)):
            set_cell_margins(cell, top=70, bottom=70, start=120, end=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col in (0, 4) else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(
                run,
                size=9.5 if col != 1 else 10.5,
                color=INK,
                bold=(col == 1),
            )
    set_table_geometry(table, widths)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("오늘 헷갈린 단어:  ______________________________________________________________")
    set_run_font(run, size=8.5, color=MUTED)


def build():
    records = load_records()
    by_day = defaultdict(list)
    for record in records:
        by_day[record["day"]].append(record)

    if (
        len(records) != 480
        or set(by_day) != set(range(7, 31))
        or any(len(items) != 20 for items in by_day.values())
    ):
        raise ValueError("Expected 480 records, 20 per day from day 7 through day 30")

    doc = Document()
    configure_styles(doc)
    for section in doc.sections:
        configure_section(section)
        set_header_footer(section)

    props = doc.core_properties
    props.title = "JLPT N5·N4 일본어 통합 단어장 7–30일차"
    props.subject = "능력단어 7–30일차 통합 복습 단어장"
    props.author = ""
    props.keywords = "JLPT, 일본어, N5, N4, 단어장"

    add_cover(doc)
    for day in sorted(by_day):
        add_day_page(doc, day, by_day[day])

    doc.save(OUTPUT_PATH)
    print(f"saved {OUTPUT_PATH}")


if __name__ == "__main__":
    build()
