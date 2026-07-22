import html
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from build_vocab_book import load_records
from source_paths import KAKAO_SOURCE_DIR


BASE = Path(__file__).resolve().parent
SOURCE_DIR = KAKAO_SOURCE_DIR
OUTPUT_PATH = BASE / "basic_sentences.json"
JANOME_PATH = BASE / "tmp" / "janome"

if JANOME_PATH.exists():
    sys.path.insert(0, str(JANOME_PATH))

from janome.tokenizer import Tokenizer  # noqa: E402


JAPANESE_PATTERN = re.compile(r"[ぁ-ゖァ-ヺ一-龯々〆ヵヶ]")
HANGUL_PATTERN = re.compile(r"[가-힣]")
KANJI_PATTERN = re.compile(r"[一-龯々〆ヵヶ]+")
EQ_RUBY_PATTERN = re.compile(
    r"\\o\\ad\(\s*\\s\\up\s+\d+\((.*?)\)\s*,\s*(.*?)\)\s*$",
    re.DOTALL,
)
BLANK_OVERRIDES = {
    "d8-5": "ひと月",
    "d9-16": "みんな",
    "d9-17": "眼",
    "d10-17": "０",
    "d11-10": "いたい",
    "d11-17": "かわいい",
    "d12-8": "地下",
    "d13-6": "きれい",
    "d14-6": "ある",
    "d14-10": "いる",
    "d15-4": "掛け",
    "d15-18": {"before": "んでいます"},
    "d16-10": "出かけ",
    "d16-11": "できる",
    "d17-19": "わから",
    "d18-20": "だんだん",
    "d19-16": "もちろん",
    "d20-1": "あいさつ",
    "d20-13": "意思",
    "d20-15": "以内",
    "d20-18": "受け付け",
    "d20-19": "噓",
    "d21-9": "押し入れ",
    "d23-11": "けんか",
    "d26-20": "どろぼう",
    "d30-14": "すごい",
    "d30-19": "ひどい",
}

SOURCE_SENTENCE_OVERRIDES = {
    "d26-6": "血を流しています。",
    "d27-4": "入学試験に合格しました。",
    "d27-9": "喉が腫れて病院に行きました。",
    "d28-2": "カーテンの間から光が見えます。",
    "d28-6": "去年引っ越ししました。",
    "d29-2": "公園の真ん中に桜の木があります。",
    "d30-7": "彼の罪は重いです。",
    "d30-9": "固い食べ物は噛めません。",
    "d30-15": "素晴らしい作品を発見した。",
}

TRANSLATION_OVERRIDES = {
    "d26-8": "주의사항을 반드시 확인해 주세요.",
    "d26-16": "일기예보를 확인해 주세요.",
    "d27-3": "교통사고로 병원에 입원했습니다.",
    "d27-9": "목이 부어서 병원에 갔습니다.",
    "d27-10": "자주 이용하는 교통수단은 신칸센, 모노레일, 지하철입니다.",
    "d27-15": "그녀는 발음이 좋습니다.",
    "d27-16": "봄이 되면 꽃구경을 갑니다.",
    "d28-3": "중요한 서류는 서랍에 보관합니다.",
    "d28-4": "오랫동안 수염을 깎지 않았습니다.",
    "d28-12": "일본의 전통문화를 연구하고 있습니다.",
    "d28-15": "그의 대답을 기다리고 있습니다.",
    "d28-16": "여기가 무역센터입니다.",
    "d29-3": "이 근처에 아름다운 호수가 있습니다.",
    "d29-14": "수업 전에 반드시 예습을 해야 합니다.",
    "d30-3": "자연은 아름답습니다.",
    "d30-7": "그의 죄는 무겁습니다.",
}

TOKEN_READING_OVERRIDES = {
    ("d28-2", "間"): "あいだ",
}


def sort_key(path):
    match = re.search(r"(\d+)$", path.stem)
    return int(match.group(1)) if match else 0


def clean(text):
    return re.sub(r"\s+", " ", text).strip()


def katakana_to_hiragana(text):
    return "".join(
        chr(ord(char) - 0x60) if "ァ" <= char <= "ヶ" else char
        for char in text
    )


def tokenizer_segments(sentence, tokenizer, record_id=""):
    segments = []
    for token in tokenizer.tokenize(sentence):
        surface = token.surface
        reading = token.reading
        if reading == "*" or not KANJI_PATTERN.search(surface):
            reading = ""
        else:
            reading = katakana_to_hiragana(reading)
        reading = TOKEN_READING_OVERRIDES.get((record_id, surface), reading)
        segments.append({"text": surface, "reading": reading})
    return segments


def direct_text(element, tag_name):
    return "".join(node.text or "" for node in element.findall(qn(tag_name)))


def iter_runs(element):
    for child in element:
        if child.tag == qn("w:r"):
            yield child
        elif child.tag in {
            qn("w:hyperlink"),
            qn("w:smartTag"),
            qn("w:sdt"),
            qn("w:sdtContent"),
            qn("w:ins"),
        }:
            yield from iter_runs(child)


def parse_eq_ruby(instruction):
    match = EQ_RUBY_PATTERN.search(clean(instruction))
    if not match:
        return None
    reading = clean(match.group(1))
    base = clean(match.group(2))
    return {"text": base, "reading": reading}


def paragraph_segments(paragraph):
    segments = []
    field_instruction = None
    field_separated = False

    for run in iter_runs(paragraph._p):
        field_char = run.find(qn("w:fldChar"))
        if field_char is not None:
            field_type = field_char.get(qn("w:fldCharType"))
            if field_type == "begin":
                field_instruction = []
                field_separated = False
            elif field_type == "separate" and field_instruction is not None:
                field_separated = True
            elif field_type == "end" and field_instruction is not None:
                parsed = parse_eq_ruby("".join(field_instruction))
                if parsed:
                    segments.append(parsed)
                field_instruction = None
                field_separated = False
            continue

        instruction_text = direct_text(run, "w:instrText")
        if field_instruction is not None:
            if instruction_text and not field_separated:
                field_instruction.append(instruction_text)
            continue

        ruby = run.find(qn("w:ruby"))
        if ruby is not None:
            ruby_text = ruby.find(qn("w:rt"))
            ruby_base = ruby.find(qn("w:rubyBase"))
            reading = "" if ruby_text is None else "".join(
                node.text or "" for node in ruby_text.iter(qn("w:t"))
            )
            base = "" if ruby_base is None else "".join(
                node.text or "" for node in ruby_base.iter(qn("w:t"))
            )
            if base:
                segments.append({"text": clean(base), "reading": clean(reading)})
            continue

        text = direct_text(run, "w:t")
        if text:
            segments.append({"text": text, "reading": ""})
        if run.find(qn("w:tab")) is not None:
            segments.append({"text": "\t", "reading": ""})
        if run.find(qn("w:br")) is not None:
            segments.append({"text": "\n", "reading": ""})

    return segments


def candidate_forms(*words):
    forms = set()
    for word in words:
        if not word:
            continue
        normalized = clean(word).replace("・", "")
        pieces = re.split(r"[／/,、]", normalized)
        for piece in [normalized, *pieces]:
            piece = piece.strip()
            if not piece:
                continue
            forms.add(piece)
            if piece.endswith("だ") and len(piece) > 1:
                forms.add(piece[:-1])
            without_parentheses = re.sub(r"（[^）]*）", "", piece).strip()
            if without_parentheses:
                forms.add(without_parentheses)
            for parenthetical in re.findall(r"（([^）]+)）", piece):
                parenthetical = parenthetical.strip()
                if parenthetical and parenthetical not in {"もう"}:
                    forms.add(parenthetical)
    return sorted(forms, key=lambda value: (-len(value), value))


def locate_blank(record_id, sentence, word, source_word, tokenizer):
    override = BLANK_OVERRIDES.get(record_id)
    if override:
        if isinstance(override, dict):
            position = sentence.find(override["before"])
            if position >= 0:
                return position, position, "source-override"
        else:
            position = sentence.find(override)
            if position >= 0:
                return position, position + len(override), "source-override"
        raise ValueError(f"Could not apply blank override {override!r} to {sentence!r}")

    forms = candidate_forms(word, source_word)
    form_set = set(forms)
    cursor = 0
    for token in tokenizer.tokenize(sentence):
        surface = token.surface
        start = sentence.find(surface, cursor)
        if start < 0:
            start = cursor
        end = start + len(surface)
        cursor = end
        if token.base_form in form_set:
            return start, end, "base-form"

    for form in forms:
        position = sentence.find(form)
        if position >= 0:
            return position, position + len(form), "exact"

    kanji_candidates = []
    for form in forms:
        kanji_candidates.extend(KANJI_PATTERN.findall(form))
    for kanji in sorted(set(kanji_candidates), key=lambda value: (-len(value), value)):
        position = sentence.find(kanji)
        if position >= 0:
            return position, position + len(kanji), "kanji-stem"

    raise ValueError(f"Could not locate target {word!r} in sentence {sentence!r}")


def render_furigana_with_blank(segments, blank_start, blank_end):
    rendered = []
    cursor = 0
    blank_inserted = False

    for segment in segments:
        text = segment["text"]
        start = cursor
        end = start + len(text)
        cursor = end

        if end <= blank_start or start >= blank_end:
            if not blank_inserted and start >= blank_end:
                rendered.append("___")
                blank_inserted = True
            escaped_text = html.escape(text)
            reading = segment.get("reading", "")
            if reading:
                rendered.append(
                    f"<ruby>{escaped_text}<rt>{html.escape(reading)}</rt></ruby>"
                )
            else:
                rendered.append(escaped_text)
            continue

        if not blank_inserted:
            prefix = text[: max(0, blank_start - start)]
            if prefix:
                rendered.append(html.escape(prefix))
            rendered.append("___")
            blank_inserted = True

        suffix = text[max(0, blank_end - start) :]
        if suffix:
            rendered.append(html.escape(suffix))

    if not blank_inserted:
        rendered.append("___")

    return "".join(rendered)


def slice_segments(segments, start, end):
    sliced = []
    cursor = 0
    for segment in segments:
        text = segment["text"]
        segment_start = cursor
        segment_end = segment_start + len(text)
        cursor = segment_end
        overlap_start = max(start, segment_start)
        overlap_end = min(end, segment_end)
        if overlap_start >= overlap_end:
            continue
        relative_start = overlap_start - segment_start
        relative_end = overlap_end - segment_start
        sliced_text = text[relative_start:relative_end]
        sliced.append(
            {
                "text": sliced_text,
                "reading": segment.get("reading", "")
                if relative_start == 0 and relative_end == len(text)
                else "",
            }
        )
    return sliced


def strip_segments(segments):
    text = "".join(segment["text"] for segment in segments)
    stripped = text.strip()
    if not stripped:
        return "", []
    start = len(text) - len(text.lstrip())
    end = start + len(stripped)
    return stripped, slice_segments(segments, start, end)


def extract_example(cell):
    paragraphs = []
    for paragraph in cell.paragraphs:
        segments = paragraph_segments(paragraph)
        text = "".join(segment["text"] for segment in segments)
        if text:
            paragraphs.append((text, segments))

    japanese_index = next(
        (index for index, (text, _) in enumerate(paragraphs) if JAPANESE_PATTERN.search(text)),
        None,
    )
    if japanese_index is None:
        raise ValueError("Example cell has no Japanese sentence")

    mixed_text, mixed_segments = paragraphs[japanese_index]
    hangul_match = HANGUL_PATTERN.search(mixed_text)
    if hangul_match:
        sentence, segments = strip_segments(
            slice_segments(mixed_segments, 0, hangul_match.start())
        )
        translation_parts = [mixed_text[hangul_match.start() :]]
    else:
        sentence, segments = strip_segments(mixed_segments)
        translation_parts = []
    translation_parts.extend(
        text
        for text, _ in paragraphs[japanese_index + 1 :]
        if HANGUL_PATTERN.search(text)
    )
    translation = clean(" ".join(translation_parts))
    if not translation:
        raise ValueError(f"Example sentence has no Korean translation: {sentence}")
    return sentence, segments, translation


def extract_source_records():
    records = {}
    paths = [
        path
        for path in SOURCE_DIR.glob("능력단어*.docx")
        if not path.name.startswith("~$") and 7 <= sort_key(path) <= 30
    ]
    for path in sorted(paths, key=sort_key):
        day = sort_key(path)
        document = Document(path)
        for table in document.tables:
            for row in table.rows[1:]:
                if len(row.cells) < 5:
                    continue
                number = clean(row.cells[0].text)
                source_word = clean(row.cells[1].text)
                if not number.isdigit() or not source_word:
                    continue
                sentence, segments, translation = extract_example(row.cells[4])
                records[(day, int(number))] = {
                    "sourceWord": source_word,
                    "sentence": sentence,
                    "segments": segments,
                    "translation": translation,
                    "source": path.name,
                }
    return records


def main():
    curated_records = load_records()
    source_records = extract_source_records()
    tokenizer = Tokenizer()
    exported = []
    errors = []

    for record in curated_records:
        key = (record["day"], int(record["no"]))
        record_id = f"d{record['day']}-{record['no']}"
        source = source_records.get(key)
        if source is None:
            raise ValueError(f"Missing source example for day/no {key}")
        sentence = SOURCE_SENTENCE_OVERRIDES.get(record_id, source["sentence"])
        segments = (
            tokenizer_segments(sentence, tokenizer, record_id)
            if record_id in SOURCE_SENTENCE_OVERRIDES
            else source["segments"]
        )
        try:
            blank_start, blank_end, method = locate_blank(
                record_id,
                sentence,
                record["word"],
                source["sourceWord"],
                tokenizer,
            )
        except ValueError as error:
            errors.append(f"d{record['day']}-{record['no']}: {error}")
            continue
        sentence_with_blank = sentence[:blank_start] + "___" + sentence[blank_end:]
        furigana_with_blank = render_furigana_with_blank(
            segments, blank_start, blank_end
        )
        if "<ruby>" not in furigana_with_blank and KANJI_PATTERN.search(
            sentence_with_blank
        ):
            furigana_with_blank = render_furigana_with_blank(
                tokenizer_segments(sentence, tokenizer, record_id), blank_start, blank_end
            )
        exported.append(
            {
                "id": record_id,
                "word": record["word"],
                "sentence": sentence_with_blank,
                "sentenceFurigana": furigana_with_blank,
                "sentenceTranslation": TRANSLATION_OVERRIDES.get(
                    record_id, source["translation"]
                ),
                "blankSurface": sentence[blank_start:blank_end],
                "matchMethod": method,
                "source": source["source"],
            }
        )

    if errors:
        raise ValueError("\n".join(errors))
    if len(exported) != 480 or len({item["id"] for item in exported}) != 480:
        raise ValueError(f"Expected 480 unique sentence records, got {len(exported)}")
    if not all(item["sentence"].count("___") == 1 for item in exported):
        raise ValueError("Each sentence must contain exactly one blank")
    if not all(item["sentenceFurigana"].count("___") == 1 for item in exported):
        raise ValueError("Each furigana sentence must contain exactly one blank")
    if not all(item["sentenceTranslation"] for item in exported):
        raise ValueError("Each sentence must contain a Korean translation")

    OUTPUT_PATH.write_text(
        json.dumps(exported, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    methods = {}
    for item in exported:
        methods[item["matchMethod"]] = methods.get(item["matchMethod"], 0) + 1
    print(f"extracted {len(exported)} sentence records to {OUTPUT_PATH.name}")
    print("match methods:", methods)
    for item in exported[:5]:
        print(item["id"], item["sentence"], "=>", item["sentenceTranslation"])


if __name__ == "__main__":
    main()
