from pathlib import Path
from docx import Document

from source_paths import KAKAO_SOURCE_DIR

ROOT = KAKAO_SOURCE_DIR

def sort_key(path: Path):
    stem = path.stem
    number = "".join(ch for ch in stem if ch.isdigit())
    return int(number) if number else 0

for path in sorted(ROOT.glob("능력단어*.docx"), key=sort_key):
    doc = Document(path)
    print(f"\n### {path.name}: paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
    for i, p in enumerate(doc.paragraphs[:12]):
        text = p.text.strip()
        if text:
            print(f"P{i}: {text!r}")
    for ti, table in enumerate(doc.tables):
        print(f"TABLE {ti}: rows={len(table.rows)} cols={len(table.columns)}")
        for row in table.rows[:6]:
            print(" | ".join(cell.text.replace("\n", " / ").strip() for cell in row.cells))
